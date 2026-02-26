from __future__ import annotations
import io
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import typing
from typing import Dict, List, Any, Optional
from pathlib import Path
from datetime import datetime
import plotly.express as px
import os


class SMSPDF(FPDF):
    """Custom FPDF class with SMS branding header and footer"""
    def __init__(self, customer_name="Customer", *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.customer_name = customer_name
        self.logo_path = "assets/logo.png"

    def header(self):
        # Logo in the center
        if os.path.exists(self.logo_path):
            page_width = 210
            logo_width = 35
            x = (page_width - logo_width) / 2
            self.image(self.logo_path, x, 7, logo_width)
        
        # Left side: Company Name and Date/Time
        self.set_y(8)
        self.set_x(15)
        self.set_font("helvetica", 'B', 10)
        self.cell(0, 5, self.customer_name, ln=1, align='L')
        self.set_x(15)
        self.set_font("helvetica", '', 8)
        self.cell(0, 5, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=0, align='L')
        
        # Right side: For internal use only
        self.set_y(8)
        self.set_x(-70)
        self.set_font("helvetica", 'I', 8)
        self.cell(55, 5, "For internal use only", ln=1, align='R')
        
        # Horizontal line
        self.line(10, 22, 200, 22)
        # Add some space after the header
        self.ln(20)

    def footer(self):
        self.set_y(-15)
        self.set_font("helvetica", 'I', 8)
        # Page | X on the right
        self.cell(0, 10, f'Page | {self.page_no()}', 0, 0, 'R')


class ExportService:
    """Service for exporting customer profiles to PDF/DOCX with enhanced SMS branding"""
    
    def __init__(self):
        self.assets_dir = Path("assets")
        self.temp_dir = Path("temp")
        self.temp_dir.mkdir(exist_ok=True)

    def generate_docx(self, profile: Dict, customer_name: str) -> io.BytesIO:
        """Generate a DOCX file from customer profile"""
        doc = Document()
        
        # --- Corporate Header ---
        section = doc.sections[0]
        header = section.header
        htable = header.add_table(1, 3, width=Inches(6))
        htable.style = 'Table Grid' # Change to no border later
        
        # Left: Header text
        p_left = htable.cell(0, 0).paragraphs[0]
        r1 = p_left.add_run("Integrated Performance Assessment")
        r1.bold = True
        r1.font.size = Pt(9)
        p_left.add_run(f"\n{customer_name}").font.size = Pt(9)
        
        # Center: Logo
        if os.path.exists("assets/logo.png"):
            p_mid = htable.cell(0, 1).paragraphs[0]
            p_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_logo = p_mid.add_run()
            run_logo.add_picture("assets/logo.png", width=Inches(1.2))
            
        # Right: Revision and Page
        p_right = htable.cell(0, 2).paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_right.add_run("Report Revision 01").font.size = Pt(9)
        p_right.add_run("\nPage | ").font.size = Pt(9)
        # Note: Adding dynamic page numbers to DOCX headers is tricky with python-docx,
        # usually requires complex XML. Skipping for now or using a placeholder.
        
        doc.add_paragraph() # Spacing
        
        # --- Content ---
        title = doc.add_heading(f'Customer Profile: {customer_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Basic Data
        doc.add_heading('Basic Data', level=1)
        basic_data = profile.get('basic_data', {})
        
        self._add_field(doc, 'Company Name', basic_data.get('name', 'N/A'))
        self._add_field(doc, 'Ownership History', basic_data.get('ownership_history', 'N/A'))
        self._add_field(doc, 'Headquarters Address', basic_data.get('hq_address', 'N/A'))
        self._add_field(doc, 'Owner/Parent Company', basic_data.get('owner', 'N/A'))
        self._add_field(doc, 'Management/CEO', f"{basic_data.get('management', 'N/A')} / {basic_data.get('ceo', 'N/A')}")
        self._add_field(doc, 'Employees (FTE)', basic_data.get('fte', 'N/A'))
        self._add_field(doc, 'Financial Status', basic_data.get('financials', 'N/A'))
        self._add_field(doc, 'Recent News & Facts', basic_data.get('recent_facts', 'N/A'))
        
        # Map Section
        map_path = self._generate_static_map(profile, customer_name)
        if map_path:
            doc.add_heading('Geographic Distribution', level=1)
            doc.add_picture(str(map_path), width=Inches(6.0))
            doc.add_paragraph(f"Geographic plant distribution for {customer_name}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Locations & Installed Base
        doc.add_heading('Locations & Installed Base', level=1)
        locations = profile.get('locations', [])
        for i, loc in enumerate(locations, 1):
            doc.add_heading(f'Location {i}: {loc.get("city", "Unknown")}', level=2)
            self._add_field(doc, 'Address', loc.get('address', 'N/A'))
            
            eq_list = loc.get('installed_base', [])
            if eq_list:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Equipment'
                hdr_cells[1].text = 'OEM'
                hdr_cells[2].text = 'Year'
                hdr_cells[3].text = 'Status'
                
                for eq in eq_list:
                    row_cells = table.add_row().cells
                    if isinstance(eq, dict):
                        row_cells[0].text = str(eq.get('equipment_type', 'N/A'))
                        row_cells[1].text = str(eq.get('manufacturer', 'N/A'))
                        row_cells[2].text = str(eq.get('year_of_startup', 'N/A'))
                        row_cells[3].text = str(eq.get('status', 'N/A'))
                    else:
                        row_cells[0].text = str(eq)
                        row_cells[1].text = "N/A" # OEM
                        row_cells[2].text = "N/A" # Year
                        row_cells[3].text = "Active" # Status
            
            self._add_field(doc, 'Production Capacity', loc.get('tons_per_year', 'N/A'))
        
        # Financial History Chart
        fin_chart_path = self._generate_financial_chart(profile, customer_name)
        if fin_chart_path:
            doc.add_heading('Financial Performance (10-Year History)', level=1)
            doc.add_picture(str(fin_chart_path), width=Inches(6.0))
            doc.add_paragraph(f"Revenue and EBITDA trends for {customer_name}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Balance Sheet
        balance_sheet = profile.get('latest_balance_sheet', {})
        if balance_sheet and any(balance_sheet.values()):
            doc.add_heading('Latest Balance Sheet Summary', level=1)
            self._add_field(doc, 'Assets', balance_sheet.get('assets', 'N/A'))
            self._add_field(doc, 'Liabilities', balance_sheet.get('liabilities', 'N/A'))
            self._add_field(doc, 'Equity', balance_sheet.get('equity', 'N/A'))
        
        # Remaining Sections...
        for sec_name, sec_key in [("Market Context", "context"), ("Metallurgical Insights", "metallurgical_insights"), ("Strategic Sales Strategy", "sales_strategy")]:
            doc.add_heading(sec_name, level=1)
            sec_data = profile.get(sec_key, {})
            for key, val in sec_data.items():
                self._add_field(doc, key.replace('_', ' ').title(), str(val))

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def generate_pdf(self, profile: Dict, customer_name: str) -> io.BytesIO:
        """Generate a PDF file from customer profile"""
        pdf = SMSPDF(customer_name=customer_name)
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # No Title (moved to header)
        pdf.ln(5)
        
        # Basic Data
        self._pdf_section_header(pdf, "Basic Data")
        basic = profile.get('basic_data', {})
        self._pdf_add_field_table(pdf, "Company Name", basic.get('name', 'N/A'))
        self._pdf_add_field_table(pdf, "Ownership History", basic.get('ownership_history', 'N/A'))
        self._pdf_add_field_table(pdf, "HQ Address", basic.get('hq_address', 'N/A'))
        self._pdf_add_field_table(pdf, "CEO", basic.get('ceo', 'N/A'))
        self._pdf_add_field_table(pdf, "Employees (FTE)", basic.get('fte', 'N/A'))
        self._pdf_add_field_table(pdf, "Recent Facts", basic.get('recent_facts', 'N/A'))

        # Map
        map_path = self._generate_static_map(profile, customer_name)
        if map_path:
            pdf.ln(5)
            # Add map image, scaled to width
            pdf.image(str(map_path), x=20, w=170)
            pdf.set_font("helvetica", 'I', 8)
            pdf.cell(0, 5, f"Figure: Geographic distribution for {customer_name}", ln=True, align='C')
            pdf.ln(5)

        # Financial History Chart
        fin_chart_path = self._generate_financial_chart(profile, customer_name)
        if fin_chart_path:
            self._pdf_section_header(pdf, "Financial Performance (10-Year History)")
            pdf.image(str(fin_chart_path), x=20, w=170)
            pdf.set_font("helvetica", 'I', 8)
            pdf.cell(0, 5, f"Figure: Revenue and EBITDA trends for {customer_name}", ln=True, align='C')
            pdf.ln(5)
        
        # Balance Sheet
        balance_sheet = profile.get('latest_balance_sheet', {})
        if balance_sheet and any(balance_sheet.values()):
            self._pdf_section_header(pdf, "Latest Balance Sheet Summary")
            self._pdf_add_field_table(pdf, "Assets", balance_sheet.get('assets', 'N/A'))
            self._pdf_add_field_table(pdf, "Liabilities", balance_sheet.get('liabilities', 'N/A'))
            self._pdf_add_field_table(pdf, "Equity", balance_sheet.get('equity', 'N/A'))
        
        # Other Sections
        sections = [
            ("Locations & Installed Base", profile.get('locations', [])),
            ("Project History", profile.get('history', {})),
            ("Market Context", profile.get('context', {})),
            ("Technical Insights", profile.get('metallurgical_insights', {})),
            ("Sales Strategy", profile.get('sales_strategy', {}))
        ]

        for title, data in sections:
            self._pdf_section_header(pdf, title)
            if isinstance(data, list): # Locations
                for i, loc in enumerate(data, 1):
                    pdf.set_font("helvetica", 'B', 12)
                    pdf.cell(0, 8, self._clean_text(f"Location {i}: {loc.get('city', 'Unknown')}"), ln=True)
                    self._pdf_add_field_table(pdf, "Address", loc.get('address', 'N/A'))
                    # Equipment table
                    eq_list = loc.get('installed_base', [])
                    if eq_list:
                        self._pdf_add_equipment_table(pdf, eq_list)
            else:
                for k, v in data.items():
                    if v and str(v).lower() != 'none':
                        self._pdf_add_field_table(pdf, k.replace('_', ' ').title(), v)

        pdf_bytes = pdf.output()
        if isinstance(pdf_bytes, str):
            pdf_bytes = pdf_bytes.encode('latin-1')
        return io.BytesIO(pdf_bytes)

    def _generate_static_map(self, profile: Dict, customer_name: str) -> typing.Optional[Path]:
        """Generate a static map image of customer locations using Plotly & Kaleido"""
        try:
            locations = profile.get('locations', [])
            basic = profile.get('basic_data', {})
            
            lats = []
            lons = []
            names = []
            
            # Add HQ if available
            if basic.get('latitude') and basic.get('longitude'):
                lats.append(float(basic['latitude']))
                lons.append(float(basic['longitude']))
                names.append(f"{customer_name} HQ")
                
            # Add plant locations
            for i, loc in enumerate(locations, 1):
                if loc.get('latitude') and loc.get('longitude'):
                    lats.append(float(loc['latitude']))
                    lons.append(float(loc['longitude']))
                    names.append(f"Plant {i}: {loc.get('city', 'Location')}")
            
            if not lats:
                return None
                
            fig = px.scatter_geo(
                lat=lats, lon=lons, hover_name=names,
                projection="natural earth",
                title=f"Locations for {customer_name}"
            )
            fig.update_geos(showcountries=True, countrycolor="LightGray")
            
            file_path = self.temp_dir / f"map_{customer_name.replace(' ', '_')}.png"
            fig.write_image(str(file_path))
            return file_path
        except Exception as e:
            print(f"Error generating map: {e}")
            return None
    
    def _generate_financial_chart(self, profile: Dict, customer_name: str) -> typing.Optional[Path]:
        """Generate a financial history bar chart using Plotly"""
        try:
            fin_history = profile.get('financial_history', [])
            if not fin_history or len(fin_history) < 2:
                return None
            
            import plotly.graph_objects as go
            from plotly.subplots import make_subplots
            
            df = pd.DataFrame(fin_history)
            
            # Create figure with secondary y-axis
            fig = make_subplots(specs=[[{"secondary_y": True}]])
            
            # Add revenue bars
            fig.add_trace(
                go.Bar(name="Revenue (M EUR)", x=df['year'], y=df.get('revenue_m_eur', [0]*len(df)),
                       marker_color='lightblue'),
                secondary_y=False,
            )
            
            # Add EBITDA bars
            fig.add_trace(
                go.Bar(name="EBITDA (M EUR)", x=df['year'], y=df.get('ebitda_m_eur', [0]*len(df)),
                       marker_color='darkblue'),
                secondary_y=False,
            )
            
            fig.update_layout(
                title=f"Financial Performance - {customer_name}",
                xaxis_title="Year",
                yaxis_title="Million EUR",
                barmode='group',
                height=400,
                showlegend=True
            )
            
            file_path = self.temp_dir / f"financials_{customer_name.replace(' ', '_')}.png"
            fig.write_image(str(file_path))
            return file_path
        except Exception as e:
            print(f"Error generating financial chart: {e}")
            return None

    def _pdf_section_header(self, pdf, title):
        pdf.ln(5)
        pdf.set_font("helvetica", 'B', 14)
        pdf.set_fill_color(230, 230, 230)
        pdf.cell(0, 10, self._clean_text(title), ln=True, fill=True)
        pdf.ln(2)

    def _pdf_add_field_table(self, pdf, label, value):
        """Add field in table format (no bullet points)"""
        pdf.set_font("helvetica", 'B', 10)
        pdf.cell(60, 6, self._clean_text(label) + ":", border=0)
        pdf.set_font("helvetica", '', 10)
        pdf.multi_cell(0, 6, self._clean_text(str(value)), border=0)
        pdf.ln(2)
    
    def _pdf_add_equipment_table(self, pdf, eq_list):
        """Add equipment as a proper table in PDF"""
        pdf.set_font("helvetica", 'B', 9)
        pdf.ln(3)
        # Table headers
        pdf.cell(70, 6, "Equipment", border=1, align='C')
        pdf.cell(40, 6, "OEM", border=1, align='C')
        pdf.cell(30, 6, "Year", border=1, align='C')
        pdf.cell(40, 6, "Status", border=1, align='C')
        pdf.ln()
        
        pdf.set_font("helvetica", '', 8)
        for eq in eq_list:
            if isinstance(eq, dict):
                pdf.cell(70, 6, self._clean_text(str(eq.get('equipment_type', 'N/A'))[:30]), border=1)
                pdf.cell(40, 6, self._clean_text(str(eq.get('manufacturer', 'N/A'))[:18]), border=1)
                pdf.cell(30, 6, str(eq.get('year_of_startup', 'N/A')), border=1, align='C')
                pdf.cell(40, 6, self._clean_text(str(eq.get('status', 'N/A'))[:18]), border=1)
            else:
                pdf.cell(70, 6, self._clean_text(str(eq)[:30]), border=1)
                pdf.cell(40, 6, "N/A", border=1)
                pdf.cell(30, 6, "N/A", border=1, align='C')
                pdf.cell(40, 6, "Active", border=1)
            pdf.ln()
        pdf.ln(3)

    def _clean_text(self, text: str) -> str:
        if not text: return ""
        repl = {'€': 'EUR', '…': '...', '–': '-', '—': '-', '“': '"', '”': '"', '‘': "'", '’': "'"}
        for c, r in repl.items(): text = text.replace(c, r)
        try: return text.encode('latin-1', 'replace').decode('latin-1')
        except: return text.encode('ascii', 'ignore').decode('ascii')

    def _add_field(self, doc, label: str, value: Any):
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(str(value))

    def generate_filename(self, customer_name: str, extension: str = 'docx') -> str:
        clean = "".join(c for c in customer_name if c.isalnum() or c in (' ', '-', '_')).replace(' ', '_')
        return f"Customer_Profile_{clean}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{extension}"


export_service = ExportService()
